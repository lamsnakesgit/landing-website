import os
import json
from docx import Document
from docx.shared import Inches
from playwright.sync_api import sync_playwright

def generate_word_log():
    doc = Document()
    doc.add_heading('Логи парсинга Судебного Кабинета', 0)
    
    years = ["2026", "2025"]
    os.makedirs("output/screenshots", exist_ok=True)
    
    with sync_playwright() as p:
        browser = p.chromium.launch(headless=True)
        page = browser.new_page()
        
        for year in years:
            doc.add_heading(f'Год: {year}', level=1)
            year_dir = f"output/{year}"
            if not os.path.exists(year_dir):
                continue
                
            for filename in os.listdir(year_dir):
                if not filename.endswith(".json"): continue
                
                doc.add_heading(f'Файл: {filename}', level=2)
                filepath = os.path.join(year_dir, filename)
                
                try:
                    with open(filepath, "r", encoding="utf-8") as f:
                        cases = json.load(f)[:3] # Берем 3 дела для отчета
                except:
                    continue
                    
                for idx, case in enumerate(cases, 1):
                    data = case
                    if not isinstance(data, list) or len(data) == 0: continue
                    
                    case_num = str(data[0]).replace("/", "_")
                    doc.add_heading(f'Дело: {case_num}', level=3)
                    
                    # Пишем текст
                    doc.add_paragraph(f"Номер дела: {data[0] if len(data)>0 else '-'}")
                    doc.add_paragraph(f"Дата: {data[1] if len(data)>1 else '-'}")
                    doc.add_paragraph(f"Статус/Инфо: {data[2] if len(data)>2 else '-'}")
                    doc.add_paragraph("Ссылка на оригинал: https://office.sud.kz/")
                    
                    # Делаем скриншот портала (в идеале - карточки дела, но нужен Kalkan)
                    screenshot_path = f"output/screenshots/{case_num}.png"
                    try:
                        print(f"📸 Делаю скриншот для {case_num}...")
                        page.goto("https://office.sud.kz/")
                        page.screenshot(path=screenshot_path)
                        doc.add_picture(screenshot_path, width=Inches(5))
                    except Exception as e:
                        doc.add_paragraph(f"[Ошибка скриншота: {e}]")
                        
                    doc.add_paragraph("-" * 40)
                    
        browser.close()
        
    doc_path = "output/Parsing_Logs_Word.docx"
    doc.save(doc_path)
    print(f"✅ Успешно сгенерирован Word-документ: {doc_path}")

if __name__ == "__main__":
    generate_word_log()
